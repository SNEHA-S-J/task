import docx
from typing import Dict, Any, List
import re
from fuzzywuzzy import fuzz
import io

class DocumentProcessor:
    """Process and analyze ADGM corporate documents"""
    
    def __init__(self):
        self.document_types = {
            'Memorandum of Association': ['memorandum', 'moa', 'memorandum of association'],
            'Articles of Association': ['articles', 'aoa', 'articles of association'],
            'Register of Members': ['register of members', 'members register'],
            'Register of Directors': ['register of directors', 'directors register'],
            'UBO Form': ['ubo', 'ultimate beneficial owner', 'beneficial owner'],
            'Certificate of Incorporation': ['certificate of incorporation', 'incorporation certificate']
        }
    
    def process_document(self, file) -> Dict[str, Any]:
        """Process a single document and extract key information"""
        try:
            doc = docx.Document(file)
            full_text = self._extract_text(doc)
            
            return {
                'filename': file.name,
                'type': self._identify_document_type(file.name, full_text),
                'content': full_text,
                'paragraphs': [p.text for p in doc.paragraphs if p.text.strip()],
                'tables': self._extract_tables(doc),
                'word_count': len(full_text.split()),
                'metadata': {
                    'core_properties': self._get_core_properties(doc),
                    'sections': self._identify_sections(full_text)
                }
            }
        except Exception as e:
            return {'error': str(e), 'filename': file.name}
    
    def _extract_text(self, doc: docx.Document) -> str:
        """Extract all text from document"""
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def _identify_document_type(self, filename: str, content: str) -> str:
        """Identify the type of document based on filename and content"""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        for doc_type, keywords in self.document_types.items():
            for keyword in keywords:
                if keyword in filename_lower or keyword in content_lower:
                    return doc_type
        
        # Use fuzzy matching for better accuracy
        for doc_type, keywords in self.document_types.items():
            for keyword in keywords:
                if fuzz.partial_ratio(keyword, content_lower) > 80:
                    return doc_type
        
        return "Unknown Document Type"
    
    def _extract_tables(self, doc: docx.Document) -> List[List[str]]:
        """Extract tables from document"""
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables
    
    def _get_core_properties(self, doc: docx.Document) -> Dict[str, str]:
        """Get document properties"""
        props = {}
        if doc.core_properties:
            props['title'] = doc.core_properties.title or ''
            props['author'] = doc.core_properties.author or ''
            props['created'] = str(doc.core_properties.created) if doc.core_properties.created else ''
            props['modified'] = str(doc.core_properties.modified) if doc.core_properties.modified else ''
        return props
    
    def _identify_sections(self, content: str) -> List[str]:
        """Identify major sections in the document"""
        sections = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if line and (line.isupper() or line.endswith(':') or re.match(r'^\d+\.', line)):
                sections.append(line)
        
        return sections[:10]  # Return first 10 sections
    
    def extract_key_clauses(self, content: str, clause_types: List[str]) -> Dict[str, str]:
        """Extract specific types of clauses from document"""
        clauses = {}
        content_lower = content.lower()
        
        for clause_type in clause_types:
            # Look for clause patterns
            pattern = rf"{clause_type.lower()}.*?(?=\n\n|\n[A-Z]|$)"
            matches = re.findall(pattern, content_lower, re.DOTALL)
            if matches:
                clauses[clause_type] = matches[0].strip()
        
        return clauses
    
    def validate_document_completeness(self, doc_info: Dict[str, Any], required_elements: List[str]) -> Dict[str, Any]:
        """Check if document contains all required elements"""
        content = doc_info['content'].lower()
        
        missing_elements = []
        present_elements = []
        
        for element in required_elements:
            if element.lower() in content:
                present_elements.append(element)
            else:
                missing_elements.append(element)
        
        return {
            'missing_elements': missing_elements,
            'present_elements': present_elements,
            'completeness_score': len(present_elements) / len(required_elements) * 100
        }
